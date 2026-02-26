# backend/main.py
from fastapi import FastAPI, UploadFile, File, HTTPException, Response
import fitz
import easyocr
import cv2
import numpy as np
import tempfile
import os
from pdf2docx import Converter
import pdfplumber
import pandas as pd
from docx import Document

app = FastAPI(title="Internal Document Tools")

print("Loading EasyOCR model...")
reader = easyocr.Reader(['th', 'en'], gpu=False)
print("Model loaded!")

def process_pdf(file_bytes: bytes) -> str:
    # read PDF from direct Memory (faster than save on flie)
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    text = ""
    for page in doc:
        text += page.get_text("text") + "\n"
    return text.strip()

def process_image(file_bytes: bytes) -> str:
    # change bytes to picture format for OpenCV
    nparr = np.frombuffer(file_bytes, np.uint8)
    img = cv2.imdecode(nparr, cv2.IMREAD_COLOR)
    
    # Preprocessing  *Grayscale + CLAHE + Denoise*
    gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
    clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8,8))
    contrast_img = clahe.apply(gray)
    denoised = cv2.fastNlMeansDenoising(contrast_img, None, h=10, templateWindowSize=7, searchWindowSize=21)
    kernel = np.array([[0, -1, 0], [-1, 5,-1], [0, -1, 0]])
    sharpened = cv2.filter2D(denoised, -1, kernel)
    
    # READ TEXT
    result = reader.readtext(sharpened, detail=0, paragraph=True)
    return "\n".join(result)

@app.post("/extract")
async def extract_text(file: UploadFile = File(...)):
    # Read Bytes flies from User 
    file_bytes = await file.read()
    filename = file.filename.lower()
    
    try:     #checking files format
        if filename.endswith(".pdf"): 
            extracted_text = process_pdf(file_bytes)
        elif filename.endswith((".png", ".jpg", ".jpeg")):
            extracted_text = process_image(file_bytes)
        else:
            raise HTTPException(status_code=400, detail="PDF, PNG, JPG only")
            
        return {"filename": file.filename, "extracted_text": extracted_text}
    
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"processing error: {str(e)}")
    pass

# ================ PDF To Word ===================
@app.post("/convert/pdf-to-word")
async def convert_pdf_to_word(file: UploadFile = File(...)):
    if not file.filename.lower().endswith(".pdf"):
        raise HTTPException(status_code=400, detail="รองรับเฉพาะไฟล์ PDF")
    
    # mkTemp for PDF
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_pdf:
        tmp_pdf.write(await file.read())
        pdf_path = tmp_pdf.name
        
    docx_path = pdf_path.replace(".pdf", ".docx")
    
    try:
        cv = Converter(pdf_path)
        cv.convert(docx_path)
        cv.close()
        
        with open(docx_path, "rb") as f:
            docx_bytes = f.read()
            
        return Response(content=docx_bytes, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    finally:
        # clearTemp after done
        if os.path.exists(pdf_path): os.remove(pdf_path)
        if os.path.exists(docx_path): os.remove(docx_path)

# ==================PDF To Excel=================
@app.post("/convert/pdf-to-excel")
async def convert_pdf_to_excel(file: UploadFile = File(...)):
    if not file.filename.lower().endswith(".pdf"):
        raise HTTPException(status_code=400, detail="รองรับเฉพาะไฟล์ PDF")

    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_pdf:
        tmp_pdf.write(await file.read())
        pdf_path = tmp_pdf.name
        
    xlsx_path = pdf_path.replace(".pdf", ".xlsx")
    
    try:
        all_tables = []
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                tables = page.extract_tables()
                for table in tables:
                    # transfer Table to DataFrame
                    df = pd.DataFrame(table[1:], columns=table[0])
                    all_tables.append(df)
        
        if not all_tables:
            raise HTTPException(status_code=404, detail="ไม่พบตารางในไฟล์ PDF นี้")
            
        # save all table in Excel (seperate Sheet)
        with pd.ExcelWriter(xlsx_path, engine='openpyxl') as writer:
            for i, df in enumerate(all_tables):
                df.to_excel(writer, sheet_name=f"Table_{i+1}", index=False)
                
        with open(xlsx_path, "rb") as f:
            xlsx_bytes = f.read()
            
        return Response(content=xlsx_bytes, media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
    finally:
        if os.path.exists(pdf_path): os.remove(pdf_path)
        if os.path.exists(xlsx_path): os.remove(xlsx_path)
    
# =================Image To Word===================
@app.post("/convert/image-to-word")
async def convert_image_to_word(file: UploadFile = File(...)):
    if not file.filename.lower().endswith((".png", ".jpg", ".jpeg")):
        raise HTTPException(status_code=400, detail="รองรับเฉพาะไฟล์รูปภาพ (PNG, JPG)")
        
    try:
        # ใช้ EasyOCR ดึงข้อความ (เรียกใช้ฟังก์ชันเดิมของคุณ)
        file_bytes = await file.read()
        extracted_text = process_image(file_bytes)
        
        # mk new Word and inform text
        doc = Document()
        doc.add_heading('Extract from picture', 0)
        doc.add_paragraph(extracted_text)
        
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_docx:
            docx_path = tmp_docx.name
            doc.save(docx_path)
            
        with open(docx_path, "rb") as f:
            docx_bytes = f.read()
            
        return Response(content=docx_bytes, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    finally:
        if 'docx_path' in locals() and os.path.exists(docx_path): 
            os.remove(docx_path)